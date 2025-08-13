from PIL import Image
import numpy as np
from paddleocr import PaddleOCR
import pandas as pd
import unicodedata
from jiwer import wer
#import Levenshtein
import docx  # pip install python-docx


def ComparerOriginal_GT(img, ground_truth_path):
    """
    Effectue l'OCR sur une image avec PaddleOCR, compare au texte de référence
    et retourne le texte complet + métriques d'évaluation.
    
    Args:
        img: Chemin vers l'image.
        ground_truth_path: Chemin vers le fichier texte (.txt, .doc, .docx) contenant le texte attendu.
    
    Returns:
        dict: Contenant le texte OCR, le texte attendu, et les métriques (CER, WER, distance).
    """

    # --- 1. Initialisation OCR ---
    ocr = PaddleOCR(
        use_doc_orientation_classify=False,
        use_doc_unwarping=False,
        use_textline_orientation=False
    )

    # --- 2. OCR sur l'image ---
    result = ocr.predict(img)
    res = result[0]
    rec_texts = res['rec_texts']
    full_text = " ".join(map(str, rec_texts))

    # --- 3. Lire le texte ground truth selon le format ---
    if ground_truth_path.lower().endswith(".txt"):
        with open(ground_truth_path, encoding="utf-8") as f:
            ground_truth = f.read()
    elif ground_truth_path.lower().endswith(".docx"):
        doc = docx.Document(ground_truth_path)
        ground_truth = "\n".join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Format de fichier ground_truth non supporté (txt, docx)")


    return {
        'reference': full_text,
        'ground_truth': ground_truth
    }


# Fonctionne en API ne pas modifier
def recuperer_texte_ocr(img_path):
    """
    Effectue un OCR sur une image avec PaddleOCR et retourne uniquement le texte complet.

    Args:
        img_path (str): Chemin vers l'image.

    Returns:
        str: Texte complet extrait par l'OCR.
    """
    # Initialisation PaddleOCR (mode français)
    ocr = PaddleOCR(
        use_doc_orientation_classify=False,
        use_doc_unwarping=False,
        use_textline_orientation=False,
        lang='fr'
    )

    # OCR sur l'image
    result = ocr.predict(input=img_path)

    # Extraction des textes depuis le premier résultat
    rec_texts = result[0]['rec_texts']

    # Concaténation en une seule chaîne
    full_text = " ".join(map(str, rec_texts))

    return full_text

# Fonctionne en API ne pas modifier
def ComparerOriginal_GT_Normaliser(img, ground_truth_path):
    """
    Effectue l'OCR sur une image avec PaddleOCR
    et retourne uniquement le texte OCR et le texte de référence.

    Args:
        img (str): Chemin vers l'image.
        ground_truth_path (str): Chemin vers le fichier texte ou doc/docx contenant le texte attendu.

    Returns:
        dict: Contenant le texte OCR ('predicted') et le texte attendu ('reference').
    """

    # --- 1. Initialisation OCR ---
    ocr = PaddleOCR(
        use_doc_orientation_classify=False,
        use_doc_unwarping=False,
        use_textline_orientation=False
    )

    # --- 2. OCR sur l'image ---
    result = ocr.predict(img)
    res = result[0]
    rec_texts = res['rec_texts']
    full_text = " ".join(map(str, rec_texts))

    # --- 3. Lire le texte ground truth selon le format ---
    if ground_truth_path.lower().endswith(".txt"):
        with open(ground_truth_path, encoding="utf-8") as f:
            ground_truth = f.read()
    elif ground_truth_path.lower().endswith(".docx"):
        doc = docx.Document(ground_truth_path)
        ground_truth = "\n".join([p.text for p in doc.paragraphs])
    elif ground_truth_path.lower().endswith(".doc"):
        import textract  # pip install textract
        ground_truth = textract.process(ground_truth_path).decode("utf-8")
    else:
        raise ValueError("Format de fichier ground_truth non supporté (txt, doc, docx)")

    # --- 4. Normalisation ---
    def normalize_text(text):
        text = unicodedata.normalize("NFKD", text)
        text = "".join([c for c in text if not unicodedata.combining(c)])
        return text.lower().strip()

    predicted = normalize_text(full_text)
    reference = normalize_text(ground_truth)

    return {
        'reference': reference,
        'ground_truth': predicted
    }







# ==== Exemple d'utilisation ====
# if __name__ == "__main__":
#     img = r"C:\Users\Lau\Documents\Moi\1-Travail (sept 23)\3- IA\1- Formation Greta\3- Projets\16-E5GRAFANA\Berville_L_CV_IA-avril.jpg"
#     ground_truth_path = r"C:\Users\Lau\Documents\Moi\1-Travail (sept 23)\3- IA\1- Formation Greta\3- Projets\16-E5GRAFANA\Berville_L_CV_IA-avril.txt"
    
#     result = ComparerOriginal_GT(img, ground_truth_path)
#     reference = result['reference']
#     ground_truth = result['ground_truth']
#     print(reference)
#     print(ground_truth)
